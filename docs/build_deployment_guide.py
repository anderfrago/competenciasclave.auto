from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Guia_instalacion_y_despliegue.docx"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(91, 101, 114)


def font(run, size=10.5, color=None, bold=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def fixed_table(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    tbl_pr.append(indent)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(3)


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    font(r, 25, DARK_BLUE, True)
    if subtitle:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(20)
        r = p.add_run(subtitle)
        font(r, 14, MUTED)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(7 if level == 1 else 5)
    r = p.add_run(text)
    font(r, 16 if level == 1 else 13, BLUE if level == 1 else DARK_BLUE, True)


def paragraph(doc, text, bold_start=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.18
    if bold_start and text.startswith(bold_start):
        r = p.add_run(bold_start)
        font(r, bold=True)
        r = p.add_run(text[len(bold_start):])
        font(r)
    else:
        r = p.add_run(text)
        font(r)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.18
    r = p.add_run(text)
    font(r)


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.18
    r = p.add_run(text)
    font(r)


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(.18)
    p.paragraph_format.right_indent = Inches(.18)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    r.font.size = Pt(9)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    fixed_table(table, widths)
    for cell, text in zip(table.rows[0].cells, headers):
        shade(cell, "E8EEF5")
        r = cell.paragraphs[0].add_run(text)
        font(r, 10, DARK_BLUE, True)
    for values in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, values):
            r = cell.paragraphs[0].add_run(text)
            font(r, 10)
    return table


def page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def setup(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(.492)
    section.footer_distance = Inches(.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("Autopercepción de Competencias Clave · Guía de despliegue")
    font(r, 8.5, MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Uso interno · Página ")
    font(r, 8.5, MUTED)
    page_field(footer)


def build():
    doc = Document()
    setup(doc)
    add_title(doc, "Autopercepción de Competencias Clave", "Guía de instalación y despliegue en GitHub y PythonAnywhere")
    paragraph(doc, "Esta guía permite poner en marcha la aplicación Flask + Angular 21 con SQLite. Incluye la configuración de acceso con Google, verificación por correo y despliegue en PythonAnywhere.")
    add_table(doc, ["Documento", "Detalle"], [
        ("Versión", "1.0"),
        ("Fecha", date.today().strftime("%d/%m/%Y")),
        ("Destinatario", "Administración técnica del centro"),
        ("Repositorio", "Proyecto GitHub de Autopercepción de Competencias Clave"),
    ], [2700, 6660])
    add_heading(doc, "Qué incluye el proyecto")
    for item in [
        "Backend Flask con SQLite, migraciones y API protegida por roles.",
        "Frontend Angular 21 con Signals, Bootstrap y Sass.",
        "Registro por correo verificado y acceso mediante Google.",
        "Cuestionario inicial de 79 ítems, resultados históricos y rúbrica editable.",
        "Panel de tutoría y administración de cursos, tutores, competencias e ítems.",
    ]:
        bullet(doc, item)

    add_heading(doc, "1. Requisitos previos")
    paragraph(doc, "Antes de empezar, prepara una cuenta de GitHub, una cuenta de PythonAnywhere y un proyecto de Google Cloud. La dirección cc_autopercepción@gmail.com puede usarse para administrarlos, pero el nombre de usuario de PythonAnywhere lo eliges al crear la cuenta y determina la URL pública.")
    add_table(doc, ["Elemento", "Necesidad"], [
        ("GitHub", "Repositorio privado o público para alojar el código."),
        ("PythonAnywhere", "Cuenta con una versión de Python compatible con Flask 3."),
        ("Google Cloud", "Cliente OAuth 2.0 para el botón «Continuar con Google»."),
        ("SMTP", "Cuenta de correo y contraseña de aplicación para verificar cuentas locales."),
    ], [2500, 6860])
    add_heading(doc, "2. Preparar el repositorio GitHub")
    numbered(doc, "Crea un repositorio nuevo en GitHub, por ejemplo autopercepcion-competencias.")
    numbered(doc, "Añade todos los archivos del proyecto, incluido Template Formulario CCs (respuestas).xlsx. No subas el archivo .env ni la carpeta .venv.")
    numbered(doc, "Haz el primer commit y súbelo a la rama principal.")
    paragraph(doc, "El Excel se usa únicamente para cargar por primera vez competencias, ítems y textos de devolución. A partir de ese momento, la administración se realiza desde la aplicación.")

    add_heading(doc, "3. Ejecutar en local", 1)
    paragraph(doc, "En Windows, abre dos terminales en la carpeta del proyecto: una para el backend y otra para el frontend.")
    add_heading(doc, "Backend Flask", 2)
    code(doc, "python -m venv .venv")
    code(doc, ".venv\\Scripts\\python -m pip install -r backend\\requirements.txt")
    code(doc, "copy .env.example .env")
    code(doc, "cd backend && ..\\.venv\\Scripts\\python -m flask --app run.py db upgrade")
    code(doc, "..\\.venv\\Scripts\\python -m flask --app run.py init-db")
    code(doc, "..\\.venv\\Scripts\\python run.py")
    add_heading(doc, "Frontend Angular", 2)
    code(doc, "cd frontend")
    code(doc, "pnpm install")
    code(doc, "pnpm start")
    paragraph(doc, "Abre http://localhost:4200. El proxy de desarrollo deriva las llamadas /api al backend local en el puerto 5000.")

    add_heading(doc, "4. Configurar variables de entorno")
    paragraph(doc, "Copia .env.example a .env y completa las claves. Genera SECRET_KEY y JWT_SECRET_KEY con valores aleatorios largos y diferentes. Las cuentas administradoras se separan por comas en ADMIN_EMAILS.")
    add_table(doc, ["Variable", "Uso"], [
        ("ADMIN_EMAILS", "Correos con permisos de administración."),
        ("DATABASE_URL", "Ruta de SQLite. En producción usa una ruta absoluta bajo /home."),
        ("FRONTEND_URL / BACKEND_URL", "Dirección pública de la aplicación, sin barra final."),
        ("SMTP_*", "Servidor, puerto, usuario, contraseña de aplicación y remitente."),
        ("GOOGLE_CLIENT_ID / SECRET", "Credenciales OAuth del proyecto Google Cloud."),
        ("SPA_DIST_PATH", "Carpeta browser generada por el build de Angular en producción."),
    ], [2850, 6510])
    add_heading(doc, "5. Acceso con Google", 1)
    numbered(doc, "En Google Cloud Console crea un proyecto y configura la pantalla de consentimiento OAuth.")
    numbered(doc, "Crea un cliente OAuth de tipo Aplicación web.")
    numbered(doc, "Añade como origen autorizado la URL de PythonAnywhere: https://TU_USUARIO.pythonanywhere.com.")
    numbered(doc, "Añade como URI de redirección autorizado: https://TU_USUARIO.pythonanywhere.com/api/auth/google/callback.")
    numbered(doc, "Copia el identificador y el secreto en GOOGLE_CLIENT_ID y GOOGLE_CLIENT_SECRET del archivo .env.")
    paragraph(doc, "Las cuentas de Google, incluidas educación.navarra.es, se validan mediante Google. La administración puede asignar como tutor a cualquier cuenta registrada y gestionar su rol desde el panel de usuarios.")

    add_heading(doc, "6. Verificación por correo", 1)
    paragraph(doc, "Para el registro con contraseña es necesario un SMTP funcional. Con Gmail se recomienda crear una contraseña de aplicación con la verificación en dos pasos activada; no uses la contraseña habitual. Indica el remitente en SMTP_FROM y guarda la contraseña únicamente en .env.")

    add_heading(doc, "7. Desplegar en PythonAnywhere")
    paragraph(doc, "La configuración siguiente sirve para una aplicación Flask manual. Sustituye TU_USUARIO por el nombre de usuario real de PythonAnywhere y RUTA_PROYECTO por la carpeta donde hayas clonado el repositorio.")
    add_heading(doc, "Clonar y preparar", 2)
    numbered(doc, "En la consola Bash de PythonAnywhere clona el repositorio: git clone URL_DEL_REPOSITORIO cc-autopercepcion.")
    numbered(doc, "Crea el entorno: mkvirtualenv --python=/usr/bin/python3.13 cc-autopercepcion-env. Si esa versión no está disponible, elige una compatible que ofrezca PythonAnywhere y usa la misma al crear la web app.")
    numbered(doc, "Instala las dependencias: pip install -r ~/cc-autopercepcion/backend/requirements.txt.")
    numbered(doc, "Construye Angular en tu equipo con pnpm build y sube o versiona la carpeta frontend/dist/autopercepcion-cc/browser. Alternativamente instala Node compatible en PythonAnywhere y ejecuta pnpm install && pnpm build allí.")
    add_heading(doc, "Archivo .env de producción", 2)
    paragraph(doc, "Crea ~/cc-autopercepcion/.env sin subirlo a GitHub. Como DATABASE_URL usa sqlite:////home/TU_USUARIO/cc-autopercepcion/backend/instance/autopercepcion.db. Para FRONTEND_URL y BACKEND_URL utiliza https://TU_USUARIO.pythonanywhere.com. Define SPA_DIST_PATH=/home/TU_USUARIO/cc-autopercepcion/frontend/dist/autopercepcion-cc/browser.")
    add_heading(doc, "Crear la aplicación web", 2)
    numbered(doc, "En la pestaña Web crea una nueva web app, selecciona Manual configuration y la misma versión de Python usada en el entorno virtual.")
    numbered(doc, "Asigna el virtualenv /home/TU_USUARIO/.virtualenvs/cc-autopercepcion-env.")
    numbered(doc, "Abre el archivo WSGI indicado por PythonAnywhere y reemplaza su contenido por el ejemplo siguiente, ajustando el usuario.")
    code(doc, "import sys; sys.path.insert(0, '/home/TU_USUARIO/cc-autopercepcion/backend'); from run import app as application")
    numbered(doc, "Pulsa Reload. Después abre https://TU_USUARIO.pythonanywhere.com/api/health; debe devolver {\"status\":\"ok\"}.")
    paragraph(doc, "El backend sirve la compilación Angular cuando SPA_DIST_PATH apunta a la carpeta browser. Esto conserva las rutas /api para Flask y permite abrir enlaces directos como /acceso o /invitacion/CODIGO.")

    add_heading(doc, "Inicializar o actualizar la base de datos", 2)
    code(doc, "cd ~/cc-autopercepcion/backend")
    code(doc, "flask --app run.py db upgrade")
    code(doc, "flask --app run.py init-db")
    paragraph(doc, "init-db puede ejecutarse de nuevo sin duplicar las competencias. Antes de actualizar el código en producción, descarga una copia del archivo SQLite como medida de seguridad.")

    add_heading(doc, "8. Puesta en marcha y comprobación")
    add_table(doc, ["Comprobación", "Resultado esperado"], [
        ("/api/health", "Respuesta JSON con status: ok."),
        ("Crear cuenta local", "Llega un correo y el enlace permite verificar la cuenta."),
        ("Google", "El inicio redirige a Google y vuelve a la aplicación."),
        ("Curso", "Administración crea un curso y copia su enlace de invitación."),
        ("Tutor", "Una cuenta Google @cuatrovientos.org puede asignarse al curso."),
        ("Cuestionario", "Un alumno inscrito completa las 79 respuestas y ve 7 resultados."),
        ("Tutoría", "Se muestran medias, niveles, tabla individual e historial del curso."),
    ], [2700, 6660])
    add_heading(doc, "Operación habitual", 1)
    for item in [
        "Crear los cursos de cada año académico desde Administración y asignar sus tutores.",
        "Enviar a cada grupo el enlace de invitación con su código único.",
        "Permitir que el alumnado repita el formulario; el sistema conserva todos los intentos.",
        "Revisar los paneles de tutoría tomando como referencia el último intento de cada estudiante y su evolución histórica.",
        "Editar desde Administración los umbrales de cada rúbrica, los textos de devolución y el mensaje final de ánimo.",
    ]:
        bullet(doc, item)
    add_heading(doc, "Privacidad y seguridad", 1)
    paragraph(doc, "La pantalla de registro enlaza a https://cuatrovientos.org/rgpd/. Mantén el repositorio y la base de datos bajo control del centro, limita el acceso de administración y tutoría a cuentas autorizadas, y no incluyas secretos en GitHub. Las copias de SQLite deben almacenarse en una ubicación protegida.")
    add_heading(doc, "Actualizar la aplicación", 1)
    numbered(doc, "Realiza los cambios y pruebas en local.")
    numbered(doc, "Sube los cambios a GitHub.")
    numbered(doc, "En PythonAnywhere ejecuta git pull, instala nuevas dependencias si las hubiera, ejecuta flask --app run.py db upgrade, vuelve a construir Angular si cambió el frontend y pulsa Reload en la pestaña Web. Después repite las comprobaciones esenciales.")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
    print(OUTPUT)
